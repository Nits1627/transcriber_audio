import io
import json
import re
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple, Dict
from datetime import datetime

import numpy as np
import pandas as pd
import streamlit as st
from faster_whisper import WhisperModel
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import noisereduce as nr
import librosa
import soundfile as sf
from scipy import signal

# Fix for torchaudio compatibility with pyannote.audio
import torchaudio
if not hasattr(torchaudio, 'set_audio_backend'):
    # Monkey patch for newer torchaudio versions
    torchaudio.set_audio_backend = lambda x: None

# Speaker diarization imports
try:
    from pyannote.audio import Pipeline
    DIARIZATION_AVAILABLE = True
except ImportError:
    DIARIZATION_AVAILABLE = False
    Pipeline = None

# ----------------------------
# Helpers
# ----------------------------

@dataclass
class Segment:
    start: float
    end: float
    text: str
    speaker: Optional[str] = None

def sanitize_filename(name: str) -> str:
    name = re.sub(r"[^\w\-. ]+", "_", name, flags=re.UNICODE)
    name = re.sub(r"\s+", " ", name).strip()
    return name

def list_wav_files(input_dir: Path) -> List[Path]:
    """Find all WAV files in directory, excluding macOS metadata."""
    files = [p for p in input_dir.rglob("*.wav") if p.is_file()]
    # Skip macOS metadata folders/files
    files = [p for p in files if "__MACOSX" not in p.parts and not p.name.startswith("._")]
    return sorted(files)

def safe_unzip(zip_bytes: bytes, dest_dir: Path) -> None:
    """Safely extract ZIP into dest_dir (prevents zip-slip path traversal)."""
    with zipfile.ZipFile(io.BytesIO(zip_bytes)) as z:
        for member in z.infolist():
            # Skip directories and macOS junk
            if member.is_dir():
                continue
            if member.filename.startswith("__MACOSX/") or Path(member.filename).name.startswith("._"):
                continue

            target_path = dest_dir / member.filename
            # Prevent path traversal
            resolved = target_path.resolve()
            if not str(resolved).startswith(str(dest_dir.resolve())):
                raise RuntimeError("Blocked unsafe ZIP path (possible zip-slip).")

            target_path.parent.mkdir(parents=True, exist_ok=True)
            with z.open(member) as src, open(target_path, "wb") as dst:
                import shutil
                shutil.copyfileobj(src, dst)

def format_srt_time(seconds: float) -> str:
    ms = int(round(seconds * 1000.0))
    hh = ms // 3600000
    ms %= 3600000
    mm = ms // 60000
    ms %= 60000
    ss = ms // 1000
    ms %= 1000
    return f"{hh:02d}:{mm:02d}:{ss:02d},{ms:03d}"

def format_timestamp(seconds: float) -> str:
    """Format timestamp for document display."""
    total_seconds = int(seconds)
    hours = total_seconds // 3600
    minutes = (total_seconds % 3600) // 60
    secs = total_seconds % 60
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"

def enhance_audio(
    input_path: Path,
    output_path: Path,
    noise_reduce_strength: float = 0.7,
    apply_normalization: bool = True,
    apply_highpass: bool = True,
    highpass_cutoff: int = 80,
) -> None:
    """
    Enhance audio quality by reducing noise and improving clarity.
    
    Args:
        input_path: Path to input WAV file
        output_path: Path to save enhanced WAV file
        noise_reduce_strength: Strength of noise reduction (0.0-1.0)
        apply_normalization: Whether to normalize audio volume
        apply_highpass: Whether to apply high-pass filter to remove low-frequency noise
        highpass_cutoff: Cutoff frequency for high-pass filter (Hz)
    """
    # Load audio file
    audio, sr = librosa.load(str(input_path), sr=None, mono=True)
    
    # Apply high-pass filter to remove low-frequency rumble/noise (ship engines, etc.)
    if apply_highpass:
        nyquist = sr / 2
        cutoff_normalized = highpass_cutoff / nyquist
        b, a = signal.butter(4, cutoff_normalized, btype='high')
        audio = signal.filtfilt(b, a, audio)
    
    # Reduce noise using spectral gating
    # This is particularly effective for stationary noise like ship engines
    reduced_noise = nr.reduce_noise(
        y=audio,
        sr=sr,
        stationary=True,  # Good for consistent background noise
        prop_decrease=noise_reduce_strength,
        freq_mask_smooth_hz=500,  # Smooth frequency masking
        time_mask_smooth_ms=50,   # Smooth time masking
    )
    
    # Normalize audio to improve volume consistency
    if apply_normalization:
        # Peak normalization
        max_val = np.abs(reduced_noise).max()
        if max_val > 0:
            reduced_noise = reduced_noise / max_val * 0.95  # Leave some headroom
    
    # Save enhanced audio
    output_path.parent.mkdir(parents=True, exist_ok=True)
    sf.write(str(output_path), reduced_noise, sr, subtype='PCM_16')

def perform_speaker_diarization(
    audio_path: Path,
    hf_token: str,
    min_speakers: Optional[int] = None,
    max_speakers: Optional[int] = None,
) -> Dict[str, List[Tuple[float, float]]]:
    """
    Perform speaker diarization on audio file.
    
    Returns:
        Dictionary mapping speaker IDs to list of (start, end) time tuples
    """
    if not DIARIZATION_AVAILABLE:
        raise RuntimeError("pyannote.audio not installed")
    
    # Load pipeline
    pipeline = Pipeline.from_pretrained(
        "pyannote/speaker-diarization-3.1",
        use_auth_token=hf_token
    )
    
    # Run diarization
    diarization_params = {}
    if min_speakers is not None:
        diarization_params["min_speakers"] = min_speakers
    if max_speakers is not None:
        diarization_params["max_speakers"] = max_speakers
    
    diarization = pipeline(str(audio_path), **diarization_params)
    
    # Convert to dictionary format
    speaker_timeline = {}
    for turn, _, speaker in diarization.itertracks(yield_label=True):
        speaker_id = f"Speaker {speaker.split('_')[-1] if '_' in speaker else speaker}"
        if speaker_id not in speaker_timeline:
            speaker_timeline[speaker_id] = []
        speaker_timeline[speaker_id].append((turn.start, turn.end))
    
    return speaker_timeline

def align_speakers_with_transcript(
    segments: List[Segment],
    speaker_timeline: Dict[str, List[Tuple[float, float]]]
) -> List[Segment]:
    """
    Assign speaker IDs to transcript segments based on temporal overlap.
    """
    # Create flat list of (start, end, speaker) for easier lookup
    speaker_segments = []
    for speaker_id, times in speaker_timeline.items():
        for start, end in times:
            speaker_segments.append((start, end, speaker_id))
    
    # Sort by start time
    speaker_segments.sort(key=lambda x: x[0])
    
    # Assign speakers to transcript segments
    for segment in segments:
        seg_mid = (segment.start + segment.end) / 2
        
        # Find speaker with maximum overlap
        best_speaker = None
        max_overlap = 0
        
        for spk_start, spk_end, spk_id in speaker_segments:
            # Calculate overlap
            overlap_start = max(segment.start, spk_start)
            overlap_end = min(segment.end, spk_end)
            overlap = max(0, overlap_end - overlap_start)
            
            if overlap > max_overlap:
                max_overlap = overlap
                best_speaker = spk_id
        
        segment.speaker = best_speaker if best_speaker else "Unknown"
    
    return segments

def calculate_speaker_stats(
    segments: List[Segment]
) -> Dict[str, Dict[str, float]]:
    """
    Calculate speaking time statistics for each speaker.
    """
    speaker_times = {}
    total_time = 0
    
    for seg in segments:
        duration = seg.end - seg.start
        speaker = seg.speaker or "Unknown"
        
        if speaker not in speaker_times:
            speaker_times[speaker] = 0
        speaker_times[speaker] += duration
        total_time += duration
    
    # Calculate percentages
    stats = {}
    for speaker, time in speaker_times.items():
        stats[speaker] = {
            "total_time": round(time, 2),
            "percentage": round((time / total_time * 100) if total_time > 0 else 0, 1)
        }
    
    return stats

def write_txt(path: Path, segments: List[Segment], include_speakers: bool = False) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    
    if include_speakers and segments and segments[0].speaker:
        lines = []
        for seg in segments:
            speaker_label = f"[{seg.speaker}] " if seg.speaker else ""
            timestamp = f"({format_timestamp(seg.start)} - {format_timestamp(seg.end)})"
            lines.append(f"{speaker_label}{timestamp}: {seg.text}")
        content = "\n".join(lines)
    else:
        # Original format without speakers
        content = " ".join(seg.text for seg in segments)
    
    path.write_text(content.strip() + "\n", encoding="utf-8")

def write_json(path: Path, payload: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

def write_srt(path: Path, segments: List[Segment], include_speakers: bool = False) -> None:
    lines: List[str] = []
    for i, seg in enumerate(segments, start=1):
        lines.append(str(i))
        lines.append(f"{format_srt_time(seg.start)} --> {format_srt_time(seg.end)}")
        
        # Add speaker label if available
        text = seg.text.strip()
        if include_speakers and seg.speaker:
            text = f"[{seg.speaker}] {text}"
        
        lines.append(text)
        lines.append("")
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")

def transcribe_file(
    model: WhisperModel,
    wav_path: Path,
    language: Optional[str],
    task: str,
    beam_size: int,
    vad_filter: bool,
    vad_min_silence_ms: int,
) -> Tuple[str, List[Segment]]:
    segments_out: List[Segment] = []
    text_parts: List[str] = []

    vad_parameters = dict(min_silence_duration_ms=vad_min_silence_ms) if vad_filter else None

    seg_iter, _info = model.transcribe(
        str(wav_path),
        language=language,
        task=task,
        beam_size=beam_size,
        vad_filter=vad_filter,
        vad_parameters=vad_parameters,
    )

    for s in seg_iter:
        t = (s.text or "").strip()
        if not t:
            continue
        segments_out.append(Segment(start=float(s.start), end=float(s.end), text=t))
        text_parts.append(t)

    full_text = " ".join(text_parts).strip()
    return full_text, segments_out

def create_transcript_document(transcripts: List[dict], output_path: Path) -> None:
    """Create a professional Word document with all transcriptions."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Audio Transcription Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    meta_para = doc.add_paragraph()
    meta_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
    meta_para.add_run(f"Total Files: {len(transcripts)}\n").italic = True
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Process each transcript
    for idx, trans in enumerate(transcripts, 1):
        # File header
        heading = doc.add_heading(f"{idx}. {trans['filename']}", level=1)
        heading_format = heading.runs[0].font
        heading_format.color.rgb = RGBColor(0, 102, 204)
        
        # Status indicator
        status_para = doc.add_paragraph()
        status_run = status_para.add_run(f"Status: {trans['status']}")
        if trans['status'] == 'ok':
            status_run.font.color.rgb = RGBColor(0, 128, 0)
        else:
            status_run.font.color.rgb = RGBColor(255, 0, 0)
        status_run.bold = True
        
        if trans['status'] == 'ok':
            # Full transcript
            doc.add_heading('Full Transcript:', level=2)
            transcript_para = doc.add_paragraph(trans['text'])
            transcript_para.paragraph_format.line_spacing = 1.5
            
            # Timestamped segments
            if trans.get('segments'):
                doc.add_heading('Timestamped Transcript:', level=2)
                for seg in trans['segments']:
                    seg_para = doc.add_paragraph()
                    time_run = seg_para.add_run(f"[{format_timestamp(seg['start'])} - {format_timestamp(seg['end'])}] ")
                    time_run.font.color.rgb = RGBColor(128, 128, 128)
                    time_run.bold = True
                    seg_para.add_run(seg['text'])
        else:
            # Error message
            error_para = doc.add_paragraph(f"Error: {trans.get('error', 'Unknown error')}")
            error_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)
        
        # Separator
        doc.add_paragraph('_' * 80)
        doc.add_paragraph()
    
    # Save document
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

def zip_folder_bytes(folder: Path) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as z:
        for p in folder.rglob("*"):
            if p.is_file():
                z.write(p, arcname=str(p.relative_to(folder)))
    buf.seek(0)
    return buf.read()

# ----------------------------
# UI
# ----------------------------

st.set_page_config(page_title="Audio Transcriber", layout="wide", page_icon="🎙️")

# Custom CSS for better UI
st.markdown("""
    <style>
    .main-header {
        text-align: center;
        color: #1f77b4;
        padding: 20px 0;
    }
    .success-box {
        padding: 20px;
        background-color: #d4edda;
        border-left: 5px solid #28a745;
        margin: 10px 0;
    }
    .info-box {
        padding: 15px;
        background-color: #d1ecf1;
        border-left: 5px solid #0c5460;
        margin: 10px 0;
    }
    </style>
""", unsafe_allow_html=True)

st.markdown("<h1 class='main-header'>🎙️ Audio Transcription Tool</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align: center; color: #666;'>Upload WAV files individually or as a ZIP archive for automatic transcription</p>", unsafe_allow_html=True)

# Sidebar settings
with st.sidebar:
    st.header("⚙️ Settings")
    
    st.subheader("🎙️ Audio Enhancement")
    enable_enhancement = st.checkbox(
        "Enable Noise Reduction",
        value=True,
        help="Recommended for noisy environments (ships, factories, etc.)"
    )
    
    if enable_enhancement:
        noise_strength = st.slider(
            "Noise Reduction Strength",
            0.0, 1.0, 0.7, 0.1,
            help="Higher = more aggressive noise removal"
        )
        apply_highpass = st.checkbox(
            "Remove Low-Frequency Noise",
            value=True,
            help="Removes rumble from engines, ventilation, etc."
        )
        if apply_highpass:
            highpass_freq = st.slider(
                "High-Pass Filter Cutoff (Hz)",
                50, 200, 80, 10,
                help="Frequencies below this will be removed"
            )
        normalize_audio = st.checkbox(
            "Normalize Volume",
            value=True,
            help="Improves volume consistency"
        )
    
    st.divider()
    st.subheader("👥 Speaker Diarization")
    
    enable_diarization = st.checkbox(
        "Enable Speaker Identification",
        value=False,
        help="Identify and label different speakers (requires Hugging Face token)"
    )
    
    if enable_diarization:
        if not DIARIZATION_AVAILABLE:
            st.error("⚠️ pyannote.audio not installed. Install with: pip install pyannote.audio")
            enable_diarization = False
        else:
            hf_token = st.text_input(
                "Hugging Face Token",
                type="password",
                placeholder="Enter your Hugging Face token here",
                help="Required for speaker diarization. Get token at https://huggingface.co/settings/tokens"
            )
            
            col1, col2 = st.columns(2)
            with col1:
                min_speakers = st.number_input(
                    "Min Speakers",
                    min_value=1,
                    max_value=10,
                    value=None,
                    help="Minimum expected speakers (optional)"
                )
            with col2:
                max_speakers = st.number_input(
                    "Max Speakers",
                    min_value=1,
                    max_value=10,
                    value=None,
                    help="Maximum expected speakers (optional)"
                )
            
            st.info("💡 Diarization adds ~5-15 sec per minute of audio")
    
    st.divider()
    st.subheader("🤖 Transcription")
    
    model_size = st.selectbox(
        "Whisper Model",
        ["tiny", "base", "small", "medium", "large-v2", "large-v3"],
        index=3,
        help="Larger models are more accurate but slower"
    )
    
    language = st.text_input(
        "Language Code",
        value="",
        placeholder="e.g., en, es, fr (blank = auto-detect)",
        help="Leave blank for automatic language detection"
    )
    
    task = st.selectbox(
        "Task",
        ["transcribe", "translate"],
        index=0,
        help="Transcribe: keep original language | Translate: convert to English"
    )
    
    beam = st.slider("Beam Size", 1, 10, 5, help="Higher = more accurate but slower")
    vad = st.checkbox("Enable VAD Filter", value=True, help="Voice Activity Detection - removes silence")
    vad_min_silence = st.slider("VAD Min Silence (ms)", 200, 2000, 500, step=100)
    
    st.divider()
    st.info("💡 **Tip:** Use medium or large models for best accuracy")

@st.cache_resource
def load_model_cached(size: str) -> WhisperModel:
    with st.spinner(f"Loading {size} model..."):
        return WhisperModel(size, device="auto", compute_type="auto")

# Load model
model = load_model_cached(model_size)

# Upload section
st.subheader("📁 Upload Files")

upload_mode = st.radio(
    "Choose upload method:",
    ["Upload individual WAV files", "Upload ZIP archive"],
    horizontal=True
)

uploaded_files = []
zip_file = None

if upload_mode == "Upload individual WAV files":
    uploaded_files = st.file_uploader(
        "Select WAV files",
        type=["wav"],
        accept_multiple_files=True,
        help="You can select multiple files at once"
    )
    if uploaded_files:
        st.success(f"✅ {len(uploaded_files)} file(s) selected")
else:
    zip_file = st.file_uploader(
        "Select ZIP archive containing WAV files",
        type=["zip"],
        accept_multiple_files=False,
        help="ZIP file should contain WAV files"
    )
    if zip_file:
        st.success(f"✅ ZIP file uploaded: {zip_file.name}")

st.divider()

# Output format info
with st.expander("📄 Output Formats", expanded=False):
    st.markdown("""
    **For each audio file, you'll get:**
    - 📝 `.txt` - Plain text transcript
    - 🎬 `.srt` - Subtitle file with timestamps
    - 📊 `.json` - Structured data with segments
    
    **Summary files:**
    - 📋 `index.csv` - Overview of all files
    - 📄 `all_transcripts.txt` - Combined text file
    - 📘 `transcription_report.docx` - Professional Word document
    
    **Download:**
    - 🗜️ Single ZIP file with all outputs
    """)

# Start button
start = st.button("🚀 Start Transcription", type="primary", use_container_width=True)

if start:
    # Validation
    if upload_mode == "Upload individual WAV files" and not uploaded_files:
        st.error("❌ Please upload at least one WAV file")
        st.stop()
    
    if upload_mode == "Upload ZIP archive" and not zip_file:
        st.error("❌ Please upload a ZIP file")
        st.stop()
    
    language_opt = language.strip() or None
    
    with tempfile.TemporaryDirectory() as tmp:
        tmp_root = Path(tmp)
        input_dir = tmp_root / "input"
        out_root = tmp_root / "output"
        tx_dir = out_root / "transcripts"
        
        input_dir.mkdir(parents=True, exist_ok=True)
        tx_dir.mkdir(parents=True, exist_ok=True)
        
        # Prepare input files
        st.info("📦 Preparing files...")
        
        if upload_mode == "Upload ZIP archive":
            try:
                safe_unzip(zip_file.getvalue(), input_dir)
                st.success("✅ ZIP extracted successfully")
            except Exception as e:
                st.error(f"❌ Failed to extract ZIP: {e}")
                st.stop()
        else:
            # Save individual files
            for uploaded_file in uploaded_files:
                file_path = input_dir / uploaded_file.name
                file_path.write_bytes(uploaded_file.getvalue())
            st.success(f"✅ {len(uploaded_files)} file(s) ready for processing")
        
        # Find WAV files
        wav_files = list_wav_files(input_dir)
        
        if not wav_files:
            st.error("❌ No WAV files found. Please ensure your files have .wav extension")
            st.stop()
        
        st.markdown(f"<div class='info-box'>🎵 Found <b>{len(wav_files)}</b> WAV file(s) to transcribe</div>", unsafe_allow_html=True)
        
        # Progress tracking
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        rows = []
        all_text_blocks = []
        transcript_data = []
        
        # Process each file
        for idx, wav_path in enumerate(wav_files, start=1):
            rel = wav_path.relative_to(input_dir)
            base = sanitize_filename(rel.stem)
            
            status_text.markdown(f"🔄 **Processing {idx}/{len(wav_files)}:** `{rel.name}`")
            progress_bar.progress(int((idx - 1) / len(wav_files) * 100))
            
            txt_path = tx_dir / f"{base}.txt"
            srt_path = tx_dir / f"{base}.srt"
            json_path = tx_dir / f"{base}.json"
            
            try:
                # Audio enhancement
                audio_to_transcribe = wav_path
                if enable_enhancement:
                    # Save to temporary location (won't be included in output)
                    enhanced_path = wav_path.parent / f"{wav_path.stem}_enhanced.wav"
                    
                    status_text.markdown(f"🔧 **Enhancing audio:** `{rel.name}`")
                    enhance_audio(
                        input_path=wav_path,
                        output_path=enhanced_path,
                        noise_reduce_strength=noise_strength,
                        apply_normalization=normalize_audio,
                        apply_highpass=apply_highpass,
                        highpass_cutoff=highpass_freq if apply_highpass else 80,
                    )
                    audio_to_transcribe = enhanced_path
                
                # Transcribe
                status_text.markdown(f"🎤 **Transcribing:** `{rel.name}`")
                full_text, segments = transcribe_file(
                    model=model,
                    wav_path=audio_to_transcribe,
                    language=language_opt,
                    task=task,
                    beam_size=beam,
                    vad_filter=vad,
                    vad_min_silence_ms=vad_min_silence,
                )
                
                # Speaker diarization
                speaker_stats = None
                if enable_diarization and hf_token:
                    try:
                        status_text.markdown(f"👥 **Identifying speakers:** `{rel.name}`")
                        speaker_timeline = perform_speaker_diarization(
                            audio_path=audio_to_transcribe,
                            hf_token=hf_token,
                            min_speakers=int(min_speakers) if min_speakers and min_speakers > 0 else None,
                            max_speakers=int(max_speakers) if max_speakers and max_speakers > 0 else None,
                        )
                        segments = align_speakers_with_transcript(segments, speaker_timeline)
                        speaker_stats = calculate_speaker_stats(segments)
                        
                        # Update full text with speakers
                        text_parts = []
                        for seg in segments:
                            speaker_label = f"[{seg.speaker}] " if seg.speaker else ""
                            text_parts.append(f"{speaker_label}{seg.text}")
                        full_text = " ".join(text_parts)
                        
                    except Exception as e:
                        st.warning(f"⚠️ Speaker diarization failed for {rel.name}: {str(e)}")
                
                # Save outputs
                write_txt(txt_path, segments, include_speakers=enable_diarization)
                write_srt(srt_path, segments, include_speakers=enable_diarization)
                
                json_data = {
                    "source_file": str(rel),
                    "language": language_opt,
                    "task": task,
                    "segments": [{"start": s.start, "end": s.end, "text": s.text, "speaker": s.speaker} for s in segments],
                }
                if speaker_stats:
                    json_data["speakers_detected"] = len(speaker_stats)
                    json_data["speaker_stats"] = speaker_stats
                
                write_json(json_path, json_data)
                
                # Build row data
                row_data = {
                    "file": str(rel),
                    "txt": str(txt_path.relative_to(out_root)),
                    "srt": str(srt_path.relative_to(out_root)),
                    "json": str(json_path.relative_to(out_root)),
                    "preview": (full_text[:150] + ("…" if len(full_text) > 150 else "")),
                    "status": "✅ ok",
                }
                
                # Add speaker count if diarization was used
                if speaker_stats:
                    row_data["speakers"] = len(speaker_stats)
                
                rows.append(row_data)
                
                all_text_blocks.append(f"===== {rel} =====\n{full_text}\n")
                
                transcript_data.append({
                    "filename": str(rel),
                    "text": full_text,
                    "segments": [{"start": s.start, "end": s.end, "text": s.text, "speaker": s.speaker} for s in segments],
                    "status": "ok",
                    "speaker_stats": speaker_stats,
                })
                
            except Exception as e:
                error_msg = str(e)
                rows.append({
                    "file": str(rel),
                    "txt": "",
                    "srt": "",
                    "json": "",
                    "preview": "",
                    "status": f"❌ error: {error_msg}",
                })
                
                transcript_data.append({
                    "filename": str(rel),
                    "text": "",
                    "segments": [],
                    "status": "error",
                    "error": error_msg,
                })
        
        progress_bar.progress(100)
        status_text.markdown("✅ **All files processed!**")
        
        # Create summary files
        st.info("📝 Generating summary documents...")
        
        index_csv = tx_dir / "index.csv"
        all_txt = tx_dir / "all_transcripts.txt"
        docx_path = tx_dir / "transcription_report.docx"
        
        df = pd.DataFrame(rows)
        df.to_csv(index_csv, index=False, encoding="utf-8")
        all_txt.write_text("\n".join(all_text_blocks).rstrip() + "\n", encoding="utf-8")
        create_transcript_document(transcript_data, docx_path)
        
        st.success("✅ Summary documents created")
        
        # Display results
        st.subheader("📊 Results")
        st.dataframe(df, use_container_width=True)
        
        # Success summary
        success_count = len([r for r in rows if "ok" in r["status"]])
        error_count = len(rows) - success_count
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total Files", len(rows))
        with col2:
            st.metric("Successful", success_count)
        with col3:
            st.metric("Errors", error_count)
        
        # Download section
        st.divider()
        st.subheader("⬇️ Download Results")
        
        zip_bytes = zip_folder_bytes(out_root)
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.download_button(
                "📦 Download All Outputs (ZIP)",
                data=zip_bytes,
                file_name=f"transcripts_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                mime="application/zip",
                use_container_width=True,
            )
        
        with col2:
            st.download_button(
                "📘 Download Word Document",
                data=docx_path.read_bytes(),
                file_name=f"transcription_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        
        # Info about what's included
        included_items = ["📝 Transcripts (.txt, .srt, .json)", "📊 Summary files (index.csv, all_transcripts.txt)", "📘 Word document"]
        
        st.info("**Package includes:** " + " • ".join(included_items))
        st.markdown("<div class='success-box'>✅ <b>Transcription Complete!</b> Your files are ready to download.</div>", unsafe_allow_html=True)